import os
from flask import jsonify, request, send_file, flash, render_template, url_for, redirect
from docxtpl import DocxTemplate
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx.shared import Inches
from . import api_bp
from website import socketio
from datetime import datetime, timedelta
from website.models import Turnos, Espera, db, Puestos, Servicios, Usuario, Calificacion
from website.user import consultar_tabla
from sqlalchemy import text
from docx2pdf import convert
import pythoncom

def cantidad_turnos(id_servicio):
    
    turnos_asignados = Turnos.query.filter_by(id_servicio=id_servicio).all()
    turnos_espera = Espera.query.filter_by(id_servicio=id_servicio).all()
    cantidad_turnos_asignados = 0
    cantidad_turnos_espera = 0

    for turno in turnos_asignados:
        if turno.fecha == datetime.now().date():
            cantidad_turnos_asignados += 1

    for turno in turnos_espera:
        if turno.fecha_solicitud == datetime.now().date():
            cantidad_turnos_espera += 1


    cantidad_turnos = cantidad_turnos_asignados + cantidad_turnos_espera

    if cantidad_turnos == 0:
        cantidad_turnos = 1
    else:
        cantidad_turnos +=1

    return cantidad_turnos

def asignar_turno(id):
    
    if Espera.query.count() == 0:
        print("La tabla esta vacia, no hay turnos pendientes")
    else:
        #Asignacion de valores al nuevo turno
        nuevo_turno = Espera.query.first()
        id_servicio = nuevo_turno.id_servicio
        id_puesto = id
        numero_turno = nuevo_turno.id_turno
        fecha = nuevo_turno.fecha_solicitud
        estado_turno = 'Asignado'

        #Creacion del nuevo turno y guardarlo en la base de datos
        turno = Turnos(id_servicio=id_servicio, id_puesto = id_puesto, numero_turno=numero_turno, fecha=fecha, estado_turno=estado_turno)
        db.session.add(turno)
        db.session.commit()

        #Envio del mensaje al visualizador
        servicio = Servicios.query.all()
        mensaje_turnero = {'codigo':servicio[int(id_servicio) - 1].codigo, 'numero_turno':numero_turno, 'puesto':id_puesto}
        socketio.emit('turno_asignado', mensaje_turnero)

        #Cambiar el estado del puesto
        puesto = Puestos.query.get(id_puesto)
        puesto.estado = "Ocupado"
        db.session.commit()

        #Eliminar de la tabla espera
        db.session.delete(nuevo_turno)
        db.session.commit()

        turnos_restantes = Espera.query.order_by(Espera.id).all()
        for i, turno in enumerate(turnos_restantes, start=1):
            turno.id = i
        db.session.commit()

def consultar_turno(app):
    from website.models import Espera

    with app.app_context():
        try:
            turnos = Espera.query.count()
            if turnos == 0:
                return "No hay turnos pendientes"
            else:
                return "Hay turnos pendientes"
        except Exception as e:
            print(f"Error al consultar turnos: {str(e)}")

def turnos_dia():

    datos = []

    for i in range(0, 5):
        fecha_actual = datetime.now().date()

        dia_semana_actual = fecha_actual.weekday()

        # Calcular la fecha de inicio de la semana actual
        fecha_inicio_semana = fecha_actual - timedelta(days=dia_semana_actual)

        dia_a_filtrar = i  # 1 representa el martes
        dia_especifico = fecha_inicio_semana + timedelta(days=dia_a_filtrar)

        # Filtrar los datos para la semana actual
        datos.append(Turnos.query.filter(Turnos.fecha == dia_especifico).count())

    return datos

def turnos_puestos():

    datos = [0, 0 ,0]

    T_puestos = Turnos.query.filter(Turnos.fecha == datetime.now().date())

    for i in T_puestos:
        if i.estado_turno == 'Completado':
            datos[ i.id_puesto - 1 ] += 1           
        
    return datos

def liberar_turno(id):

    var = False
    turnos = Turnos.query.filter_by(id_puesto = id).all()

    if turnos:
        for turno in turnos:
            if turno.estado_turno == "Asignado":
                turno.estado_turno = 'Completado'
                db.session.commit()

def generacion_reporte(fecha_inicio = 0, fecha_fin = 0):
    # Obtener los datos del formulario
    fecha_actual = datetime.now()
    fecha = fecha_actual.strftime('%Y-%m-%d %H:%M:%S')   

    # Cargar la plantilla de Word
    if fecha_inicio == 0 or fecha_fin == 0: 
        doc = DocxTemplate('C:\\Users\\AdminDpp\\Desktop\\Reportes\\Modelo_reporte.docx')

        # Renderizar la plantilla con los datos del formulario
        turnos = Turnos.query.all()
        total_turnos = Turnos.query.count()
        total_turnos_cd = 0
        total_turnos_jfs = 0
        total_turnos_dcv = 0
        total_turnos_dfs = 0
        total_turnos_v1 = 0
        total_turnos_v2 = 0
        total_turnos_v3 = 0
        total_turnos_v4 = 0
        total_turnos_v5 = 0

        for turno in turnos:
            if turno.id_servicio == 1:
                total_turnos_cd += 1
            elif turno.id_servicio == 2:
                total_turnos_jfs += 1
            elif turno.id_servicio == 3:
                total_turnos_dcv += 1
            else:
                total_turnos_dfs += 1
        
        for turno in turnos:
            if turno.id_puesto == 1:
                total_turnos_v1 += 1
            elif turno.id_puesto == 2:
                total_turnos_v2 += 1
            elif turno.id_puesto == 3: 
                total_turnos_v3 += 1
            elif turno.id_puesto == 4:
                total_turnos_v4 += 1
            else:
                total_turnos_v5 += 1   

        context = {
            'fecha': fecha,
            'total_turnos': total_turnos,
            'total_turnos_cd' : total_turnos_cd,
            'total_turnos_jfs' : total_turnos_jfs,
            'total_turnos_dcv' : total_turnos_dcv,
            'total_turnos_dfs' : total_turnos_dfs,
            'total_turnos_v1' : total_turnos_v1,
            'total_turnos_v2' : total_turnos_v2,
            'total_turnos_v3' : total_turnos_v3,
            'total_turnos_v4' : total_turnos_v4,
            'total_turnos_v5' : total_turnos_v5
        }
    else:
        doc = DocxTemplate('C:\\Users\\AdminDpp\\Desktop\\Reportes\\Modelo_reporte_personalizado.docx')
        # Renderizar la plantilla con los datos del formulario
        turnos = Turnos.query.filter(Turnos.fecha >= fecha_inicio, Turnos.fecha <= fecha_fin).all()
        total_turnos = Turnos.query.filter(Turnos.fecha >= fecha_inicio, Turnos.fecha <= fecha_fin).count()
        total_turnos_cd = 0
        total_turnos_jfs = 0
        total_turnos_dcv = 0
        total_turnos_dfs = 0
        total_turnos_v1 = 0
        total_turnos_v2 = 0
        total_turnos_v3 = 0
        total_turnos_v4 = 0
        total_turnos_v5 = 0
        

        for turno in turnos:
            if turno.id_servicio == 1:
                total_turnos_cd += 1
            elif turno.id_servicio == 2:
                total_turnos_jfs += 1
            elif turno.id_servicio == 3:
                total_turnos_dcv += 1
            else:
                total_turnos_dfs += 1
        
        for turno in turnos:
            if turno.id_puesto == 1:
                total_turnos_v1 += 1
            elif turno.id_puesto == 2:
                total_turnos_v2 += 1
            elif turno.id_puesto == 3: 
                total_turnos_v3 += 1
            elif turno.id_puesto == 4:
                total_turnos_v4 += 1
            else:
                total_turnos_v5 += 1                 

        context = {
            'fecha': fecha,
            'fecha_inicio': fecha_inicio,
            'fecha_fin': fecha_fin,
            'total_turnos': total_turnos,
            'total_turnos_cd' : total_turnos_cd,
            'total_turnos_jfs' : total_turnos_jfs,
            'total_turnos_dcv' : total_turnos_dcv,
            'total_turnos_dfs' : total_turnos_dfs,
            'total_turnos_v1' : total_turnos_v1,
            'total_turnos_v2' : total_turnos_v2,
            'total_turnos_v3' : total_turnos_v3,
            'total_turnos_v4' : total_turnos_v4,
            'total_turnos_v5' : total_turnos_v5
        }


    servicios = ['Cambios domicilio', 'Justificaciones', 'Duplicados CV', 'Desafiliaciones']
    servicios_valores = [total_turnos_cd, total_turnos_jfs, total_turnos_dcv, total_turnos_dfs]
    colores = ['#136CB2', '#17D3E3', '#1DEEC8', '#35EE94']
    # Crear un gráfico utilizando matplotlib
    plt.bar(servicios, servicios_valores, color=colores)
    plt.xlabel('Servicios', fontweight='bold')
    plt.ylabel('Cantidad Turnos', fontweight='bold')
    plt.title('Turnos por servicio', fontweight='bold')

    # Personalizar el estilo de las barras
    plt.gca().spines['top'].set_visible(False)  # Ocultar borde superior
    plt.gca().spines['right'].set_visible(False)  # Ocultar borde derecho
    plt.gca().tick_params(axis='x', which='both', bottom=False)  # Ocultar marcas en el eje x
    plt.gca().tick_params(axis='y', which='both', left=False)  # Ocultar marcas en el eje y
    plt.grid(axis='y', linestyle='--', alpha=0.7)  # Agregar líneas de cuadrícula horizontales

    # Guardar el gráfico como una imagen
    plt.savefig('grafico_servicios.png')
    plt.clf()
    plt.close()


    puestos = ['Ventanilla 1', 'Ventanilla 2', 'Ventanilla 3', 'Ventanilla 4', 'Ventanilla 5']
    puestos_valores = [total_turnos_v1, total_turnos_v2, total_turnos_v3, total_turnos_v4, total_turnos_v5]
    colores_puestos = ['#F1F139', '#108AF0', '#F53131']
    # Crear un gráfico utilizando matplotlib
    plt.bar(puestos, puestos_valores, color=colores_puestos)
    plt.xlabel('Servicios', fontweight='bold')
    plt.ylabel('Cantidad Turnos', fontweight='bold')
    plt.title('Turnos por servicio', fontweight='bold')

    # Personalizar el estilo de las barras
    plt.gca().spines['top'].set_visible(False)  # Ocultar borde superior
    plt.gca().spines['right'].set_visible(False)  # Ocultar borde derecho
    plt.gca().tick_params(axis='x', which='both', bottom=False)  # Ocultar marcas en el eje x
    plt.gca().tick_params(axis='y', which='both', left=False)  # Ocultar marcas en el eje y
    plt.grid(axis='y', linestyle='--', alpha=0.7)  # Agregar líneas de cuadrícula horizontales

    # # Guardar el gráfico como una imagen
    plt.savefig('grafico_puestos.png')
    plt.clf()
    plt.close()

    # Insercion de graficos al documento
    doc.render(context)
    doc.add_picture('grafico_servicios.png', width=Inches(6))
    doc.add_picture('grafico_puestos.png', width=Inches(6))

    # Guardar el nuevo documento generado
    doc.save('C:\\Users\\AdminDpp\\Desktop\\Proyectos\\Turnero_CNE\\website\\static\\output.docx')

    pythoncom.CoInitialize()
    # Ruta al documento DOCX de entrada
    docx_file = "C:\\Users\\AdminDpp\\Desktop\\Proyectos\\Turnero_CNE\\website\\static\\output.docx"

    # Ruta de salida para el PDF convertido
    pdf_file = "C:\\Users\\AdminDpp\\Desktop\\Proyectos\\Turnero_CNE\\website\\static\\output.pdf"

    # Convertir el documento DOCX a PDF
    convert(docx_file, pdf_file)
    return 200


def siguiente_id_disponible():
    max_id = db.session.query(db.func.max(Usuario.id)).scalar()
    return max_id + 1 if max_id is not None else 1

def generacion_reporte_usuario(fecha_inicio = 0, fecha_fin = 0, rol = "", id_user = 0):
    # Obtener los datos del formulario
    fecha_actual = datetime.now()
    fecha = fecha_actual.strftime('%Y-%m-%d %H:%M:%S')
    user = Usuario.query.filter_by(id = id_user).first()
    nombre = user.nombre
    apellido = user.apellido
           
    # Cargar la plantilla de Word
    if fecha_inicio == 0 or fecha_fin == 0: 
        doc = DocxTemplate('C:\\Users\\AdminDpp\\Desktop\\Reportes\\Modelo_reporte _usuario.docx')
        # Renderizar la plantilla con los datos del formulario
        turnos = Turnos.query.all()
        total_turnos = 0
        total_turnos_cd = 0
        total_turnos_jfs = 0
        total_turnos_dcv = 0
        total_turnos_dfs = 0

        for turno in turnos:
            if turno.puesto.descripcion == rol:
                total_turnos += 1
                if turno.id_servicio == 1:
                    total_turnos_cd += 1
                elif turno.id_servicio == 2:
                    total_turnos_jfs += 1
                elif turno.id_servicio == 3:
                    total_turnos_dcv += 1
                else:
                    total_turnos_dfs += 1
            
        context = {
            'puesto': rol,
            'fecha': fecha,
            'nombre': nombre,
            'apellido': apellido,
            'total_turnos': total_turnos,
            'total_turnos_cd' : total_turnos_cd,
            'total_turnos_jfs' : total_turnos_jfs,
            'total_turnos_dcv' : total_turnos_dcv,
            'total_turnos_dfs' : total_turnos_dfs
        }
        
    else:
        doc = DocxTemplate('C:\\Users\\AdminDpp\\Desktop\\Reportes\\Modelo_reporte_usuario_personalizado.docx')
        # Renderizar la plantilla con los datos del formulario
        turnos = Turnos.query.filter(Turnos.fecha >= fecha_inicio, Turnos.fecha <= fecha_fin).all()
        total_turnos = 0
        total_turnos_cd = 0
        total_turnos_jfs = 0
        total_turnos_dcv = 0
        total_turnos_dfs = 0

        for turno in turnos:
            if turno.puesto.descripcion == rol:
                total_turnos += 1
                if turno.id_servicio == 1:
                    total_turnos_cd += 1
                elif turno.id_servicio == 2:
                    total_turnos_jfs += 1
                elif turno.id_servicio == 3:
                    total_turnos_dcv += 1
                else:
                    total_turnos_dfs += 1
            
        context = {
            'puesto': rol,
            'fecha': fecha,
            'nombre': nombre,
            'apellido': apellido,
            'fecha_inicio': fecha_inicio,
            'fecha_fin': fecha_fin,
            'total_turnos': total_turnos,
            'total_turnos_cd' : total_turnos_cd,
            'total_turnos_jfs' : total_turnos_jfs,
            'total_turnos_dcv' : total_turnos_dcv,
            'total_turnos_dfs' : total_turnos_dfs
        }


    servicios = ['Cambios domicilio', 'Justificaciones', 'Duplicados CV', 'Desafiliaciones']
    servicios_valores = [total_turnos_cd, total_turnos_jfs, total_turnos_dcv, total_turnos_dfs]
    colores = ['#136CB2', '#17D3E3', '#1DEEC8', '#35EE94']
    # Crear un gráfico utilizando matplotlib
    plt.bar(servicios, servicios_valores, color=colores)
    plt.xlabel('Servicios', fontweight='bold')
    plt.ylabel('Cantidad Turnos', fontweight='bold')
    plt.title('Turnos por servicio', fontweight='bold')

    # Personalizar el estilo de las barras
    plt.gca().spines['top'].set_visible(False)  # Ocultar borde superior
    plt.gca().spines['right'].set_visible(False)  # Ocultar borde derecho
    plt.gca().tick_params(axis='x', which='both', bottom=False)  # Ocultar marcas en el eje x
    plt.gca().tick_params(axis='y', which='both', left=False)  # Ocultar marcas en el eje y
    plt.grid(axis='y', linestyle='--', alpha=0.7)  # Agregar líneas de cuadrícula horizontales

    # Guardar el gráfico como una imagen
    plt.savefig('grafico_servicios.png')
    plt.clf()
    plt.close()

    # Insercion de graficos al documento
    doc.render(context)
    doc.add_picture('grafico_servicios.png', width=Inches(6))

    # Guardar el nuevo documento generado
    doc.save('C:\\Users\\AdminDpp\\Desktop\\Proyectos\\Turnero_CNE\\website\\static\\output.docx')

    pythoncom.CoInitialize()
    # Ruta al documento DOCX de entrada
    docx_file = "C:\\Users\\AdminDpp\\Desktop\\Proyectos\\Turnero_CNE\\website\\static\\output.docx"

    # Ruta de salida para el PDF convertido
    pdf_file = "C:\\Users\\AdminDpp\\Desktop\\Proyectos\\Turnero_CNE\\website\\static\\output.pdf"

    # Convertir el documento DOCX a PDF
    convert(docx_file, pdf_file)
    return 200


def calificacion_atencion(puesto, calificacion):
    
    ahora = datetime.now()
    fecha_hora_actual = ahora.strftime("%Y-%m-%d %H:%M:%S")

    ventanilla = "Ventanilla"+str(puesto)

    nueva_calificacion = Calificacion(ventanilla = ventanilla, calificacion = calificacion, fecha = fecha_hora_actual)

    db.session.add(nueva_calificacion)
    db.session.commit()
    print("Calificado correctamente")
    return 200

#---------------------------------------------RUTAS---------------------------------------------------#
    
@api_bp.route('/')
def principal():
    return "<h1> HOLA MUNDO </h1>"

@api_bp.route('/ejemplo')
def ejemplo():
    return "<h1> HOLA MUNDO COMO EJEMPLO</h1>"


@api_bp.route('/generar_turno_espera', methods=['GET'])
def generar_turno_espera():

    try:
        id_servicio = request.args.get('servicio')
        preferencial = request.args.get('preferencial')

        servicio = Servicios.query.all()

        codigo  = servicio[int(id_servicio) - 1].codigo
        ahora = datetime.now()
        fecha_hora_actual = ahora.strftime("%Y-%m-%d %H:%M:%S")

        numero_turno = cantidad_turnos(id_servicio)

        if Espera.query.count() == 0:
            query = text(f"ALTER TABLE {Espera.__tablename__} AUTO_INCREMENT = 1;")
            db.session.execute(query)
            db.session.commit()

        if preferencial == "true":

            registros_actualizar = Espera.query.order_by(Espera.id.desc()).all()
            # Actualiza los índices de los registros restantes
            for registro in registros_actualizar:
                registro.id += 1
                db.session.commit()
            
            nuevo_turno_espera = Espera(id = 1, id_turno = numero_turno, id_servicio=id_servicio, fecha_solicitud=fecha_hora_actual, codigo_turno=codigo)
            db.session.add(nuevo_turno_espera)
            db.session.commit()
        
        else:
            nuevo_turno_espera = Espera(id_turno = numero_turno, id_servicio=id_servicio, fecha_solicitud=fecha_hora_actual, codigo_turno=codigo)
            db.session.add(nuevo_turno_espera)
            db.session.commit()

        return codigo + str(numero_turno)
    
    except Exception as e:
        error = str(e)
        print(error)
        return error, 500

@api_bp.route('/datos_diarios')
def datos_diarios():
    return jsonify(turnos_dia())

@api_bp.route('/datos_puestos')
def datos_puestos():
    return jsonify(turnos_puestos())

@api_bp.route('/generar_reporte', methods=['GET'])
def generar_reporte():
    if request.method == 'GET':
        try:
            fecha_inicio = request.args.get('fecha_inicio')
            fecha_fin = request.args.get('fecha_fin')

            if len(fecha_inicio) == 0 or len(fecha_fin) == 0:
                resultado = generacion_reporte()
                return "Reporte Generado exitosamente!", resultado
            else:
                resultado = generacion_reporte(fecha_inicio, fecha_fin)
                return "Reporte Generado exitosamente!", resultado

        except Exception as e:
            error = str(e)
            print(error)
            return error, 500
        
@api_bp.route('/generar_reporte_usuario', methods=['GET'])
def generar_reporte_usuario():

    if request.method == 'GET':
        try:
            fecha_inicio = request.args.get('fecha_inicio')
            fecha_fin = request.args.get('fecha_fin')
            rol = request.args.get('rol')
            id_user = request.args.get('id_user')

            if len(fecha_inicio) == 0 or len(fecha_fin) == 0:
                resultado = generacion_reporte_usuario(0, 0, rol, id_user)
                return "Reporte Generado exitosamente!", resultado
            else:
                resultado = generacion_reporte_usuario(fecha_inicio, fecha_fin, rol, id_user)
                return "Reporte Generado exitosamente!", resultado
            

        except Exception as e:
            error = str(e)
            print(error)
            return error, 500

@api_bp.route('/crear_usuario', methods=['POST'])
def crear_usuario():
    if request.method == 'POST':
        nuevo_id = siguiente_id_disponible() 
        username = request.form['username']
        password = request.form['password']
        nombre = request.form['nombre']
        apellido = request.form['apellido']
        rol = request.form['rol']
        puesto = request.form['puesto']

        usuario_existente = Usuario.query.filter_by(username=username).first()
        if usuario_existente:
                flash("El nombre de Usuario ya existe")
                return redirect(url_for('user.user_create'))
        else:
                nuevo_usuario = Usuario(id = nuevo_id, username = username, nombre = nombre, apellido = apellido, rol = rol, puesto= puesto)

                nuevo_usuario.set_password(password)

                db.session.add(nuevo_usuario)
                db.session.commit()

                if rol == "Ventanilla":
                    puesto = Puestos.query.filter_by(descripcion = puesto).first()
                    puesto.id_user = nuevo_id
                    db.session.commit()

                flash("El Usuario fue creado correctamente", 'success')
                return redirect(url_for('user.user_management'))

@api_bp.route('/modificar_usuario', methods=['POST'])
def modificar_usuario():
    if request.method == 'POST':
        id = request.form['id']
        username = request.form['username']
        nombre = request.form['nombre']
        apellido = request.form['apellido']
        rol = request.form['rol']
        puesto = request.form['puesto']

        usuario_modificar = Usuario.query.filter_by(id = id).first()

        usuario_modificar.nombre = nombre
        usuario_modificar.apellido = apellido
        usuario_modificar.username = username
        usuario_modificar.rol = rol
        usuario_modificar.puesto = puesto
        db.session.commit()
        flash("El Usuario fue modificado correctamente", 'success')
        return redirect(url_for('user.user_management'))

@api_bp.route('/resetear_usuario', methods=['POST'])
def resetear_usuario():
    if request.method == 'POST':
        id = request.form['id']
        password = request.form['password']

        usuario_modificar = Usuario.query.filter_by(id = id).first()
        usuario_modificar.cambiar_password(password)

        flash("La contraseña fue reseteada correctamente", 'success')
        return redirect(url_for('user.user_management'))

@api_bp.route('/actualizar_tabla', methods=['GET'])
def actualizar_tabla():
    id_user = request.args.get('id_user')
    turnos = consultar_tabla(2)
    turnos_serial = []

    for turno in turnos:
        turnos_serial.append(turno.to_dict())
    
    respuesta = {
        'turnos': turnos_serial,
        'pendiente': False
    }

    asignados = Turnos.query.filter_by(estado_turno = "Asignado").all()
    
    for asignado in asignados:
        if str(asignado.puesto.id_user) == id_user:
            respuesta = {
                'turnos': turnos_serial,
                'pendiente': True
            }

    return jsonify(respuesta)

@api_bp.route('/calificar_atencion', methods=['GET'])
def calificar_atencion():
    puesto = request.args.get('puesto')
    calificacion = request.args.get('calificacion')
    resultado = calificacion_atencion(puesto, calificacion)
    return jsonify(resultado)
     
    
@socketio.on('connect')
def handle_connect():
    print('Cliente conectado.')

@socketio.on('liberar_puesto')
def handle_liberar_puesto(data):
    # Lógica para liberar el puesto y asignar un nuevo turno
    liberar_turno(data['puestoId'])
    asignar_turno(data['puestoId'])




    


 
