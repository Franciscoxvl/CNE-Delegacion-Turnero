from website import socketio
from docxtpl import DocxTemplate
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from flask import request
from docx.shared import Inches
from datetime import datetime, timedelta
from sqlalchemy import text
from website.models import Turnos, Espera, db, Servicios, Usuario, Calificacion
import subprocess
from flask_login import current_user

def consultar_numero_ventanillas(servicio):
    Usuarios_puesto = Usuario.query.filter_by(servicio = servicio).all()
    return len(Usuarios_puesto)

def obtener_puestos_ocupados(servicio):
    usuarios = Usuario.query.filter_by(servicio = servicio).all()
    return [usuario.puesto[3:] for usuario in usuarios if usuario.servicio == servicio]

def obtener_puesto_vacante(puestos_ocupados):
    puesto_num = 1
    for i in puestos_ocupados:
        puesto = f"{puesto_num}"
        if puesto not in puestos_ocupados:
            return puesto
        puesto_num += 1
    
    return 0


def cantidad_turnos(id_servicio):
    
    turnos_asignados = Turnos.query.filter_by(id_servicio=id_servicio).all()
    turnos_espera = Espera.query.filter_by(id_servicio=id_servicio).all()
    cantidad_turnos_asignados = 0
    cantidad_turnos_espera = 0

    for turno in turnos_asignados:
        if turno.fecha == datetime.now().date():
            cantidad_turnos_asignados += 1

    for turno in turnos_espera:
        if turno.fecha_solicitud == datetime.now().date():
            cantidad_turnos_espera += 1


    cantidad_turnos = cantidad_turnos_asignados + cantidad_turnos_espera

    if cantidad_turnos == 0:
        cantidad_turnos = 1
    else:
        cantidad_turnos +=1

    return cantidad_turnos

def asignar_turno(puesto):


    puesto_servicio = puesto[0:3]
    
    if Espera.query.count() == 0:
        print("La tabla esta vacia, no hay turnos pendientes")
        
    else:
        #Asignacion de valores al nuevo turno
        if puesto_servicio == "VCD":
            nuevo_turno = Espera.query.filter_by(id_servicio = 1).first()
            if nuevo_turno == None:
                mensaje_turnero = {'codigo':'N', 'numero_turno':'A', 'puesto':puesto}
                socketio.emit('turno_asignado', mensaje_turnero)
                return "Tabla vacia"
        elif puesto_servicio == "VJF":
            nuevo_turno = Espera.query.filter_by(id_servicio = 2).first()
            if nuevo_turno == None:
                mensaje_turnero = {'codigo':'N', 'numero_turno':'A', 'puesto':puesto}
                socketio.emit('turno_asignado', mensaje_turnero)
                return "Tabla vacia"
        elif puesto_servicio == "VCJ":
            nuevo_turno = Espera.query.filter_by(id_servicio = 3).first()
            if nuevo_turno == None:
                mensaje_turnero = {'codigo':'N', 'numero_turno':'A', 'puesto':puesto}
                socketio.emit('turno_asignado', mensaje_turnero)
                return "Tabla vacia"

        
        id_servicio = nuevo_turno.id_servicio
        puesto = puesto
        numero_turno = nuevo_turno.id_turno
        fecha = nuevo_turno.fecha_solicitud
        estado_turno = 'Asignado'

        #Consulta de usuario
        user = Usuario.query.filter_by(puesto = puesto).first()
        nombre = user.nombre
        apellido = user.apellido
        usuario = nombre + " " + apellido

        #Número de formulario
        if user.servicio == "Cambios de domicilio":
            numero_formulario = 1111111
        else:
            numero_formulario = 0000000

        #Creacion del nuevo turno y guardarlo en la base de datos
        turno = Turnos(id_servicio = id_servicio, puesto = puesto, usuario = usuario, numero_turno = numero_turno, numero_formulario = numero_formulario, fecha = fecha, estado_turno = estado_turno)
        db.session.add(turno)
        db.session.commit()

        #Envio del mensaje al visualizador
        servicio = Servicios.query.all()
        mensaje_turnero = {'codigo':servicio[int(id_servicio) - 1].codigo, 'numero_turno':numero_turno, 'puesto':puesto}
        socketio.emit('turno_asignado', mensaje_turnero)

        #Eliminar de la tabla espera
        db.session.delete(nuevo_turno)
        db.session.commit()

        turnos_restantes = Espera.query.order_by(Espera.id).all()
        for i, turno in enumerate(turnos_restantes, start=1):
            turno.id = i
        db.session.commit()

def consultar_turno(app):
    with app.app_context():
        try:
            turnos = Espera.query.count()
            if turnos == 0:
                mensaje = "No hay turnos pendientes"
                return mensaje
            else:
                mensaje = "Hay turnos pendientes"
                return mensaje
        except Exception as e:
            print(f"Error al consultar turnos: {str(e)}")

def turnos_dia():

    datos = []

    for i in range(0, 5):
        fecha_actual = datetime.now().date()

        dia_semana_actual = fecha_actual.weekday()

        # Calcular la fecha de inicio de la semana actual
        fecha_inicio_semana = fecha_actual - timedelta(days=dia_semana_actual)

        dia_a_filtrar = i  # 1 representa el martes
        dia_especifico = fecha_inicio_semana + timedelta(days=dia_a_filtrar)

        # Filtrar los datos para la semana actual
        datos.append(Turnos.query.filter(Turnos.fecha == dia_especifico).count())

    return datos

def generar_tespera():

    try:
        id_servicio = request.args.get('servicio')
        preferencial = request.args.get('preferencial')

        servicio = Servicios.query.all()
        
        codigo  = servicio[int(id_servicio) - 1].codigo
        ahora = datetime.now()
        fecha_hora_actual = ahora.strftime("%Y-%m-%d %H:%M:%S")

        numero_turno = cantidad_turnos(id_servicio)
        tiempo_espera_aproximado = 3
        turnos = Espera.query.filter_by(id_servicio = id_servicio).all()
        tiempo_espera = len(turnos)*tiempo_espera_aproximado
        horas, minutos = minutos_a_horas_y_minutos(tiempo_espera)
        tiempo_espera_formato = f"{ horas } hora(s) y { minutos } minutos" 

        if Espera.query.count() == 0:
            query = text(f"ALTER TABLE {Espera.__tablename__} AUTO_INCREMENT = 1;")
            db.session.execute(query)
            db.session.commit()

        if preferencial == "true":

            registros_actualizar = Espera.query.order_by(Espera.id.desc()).all()
            # Actualiza los índices de los registros restantes
            for registro in registros_actualizar:
                registro.id += 1
                db.session.commit()
            
            nuevo_turno_espera = Espera(id = 1, id_turno = numero_turno, id_servicio=id_servicio, fecha_solicitud=fecha_hora_actual, codigo_turno=codigo)
            db.session.add(nuevo_turno_espera)
            db.session.commit()
        
        else:
            nuevo_turno_espera = Espera(id_turno = numero_turno, id_servicio=id_servicio, fecha_solicitud=fecha_hora_actual, codigo_turno=codigo)
            db.session.add(nuevo_turno_espera)
            db.session.commit()

        return codigo + str(numero_turno), tiempo_espera_formato
    
    except Exception as e:
        error = str(e)
        print(error)
        return error, 500

def minutos_a_horas_y_minutos(minutos):
    horas = minutos // 60
    minutos_restantes = minutos % 60
    return horas, minutos_restantes


def turnos_puestos():

    datos = {}
    Usuarios_ventanilla_total = Usuario.query.filter_by(rol="Ventanilla").all()
    for usuario_ventanilla in Usuarios_ventanilla_total:
        datos[usuario_ventanilla.puesto] = 0


    T_puestos = Turnos.query.filter(Turnos.fecha == datetime.now().date())

    for i in T_puestos:
        if i.estado_turno == 'Completado':
            datos[i.puesto] += 1             
    return datos

def liberar_turno(puesto, n_formulario):

    var = False
    turnos = Turnos.query.filter_by(puesto = puesto).all()

    if turnos:
        for turno in turnos:
            if turno.estado_turno == "Asignado":
                turno.estado_turno = 'Completado'
                turno.numero_formulario = n_formulario
                db.session.commit()

def generate_pdf(doc_path, path):
    subprocess.call(['soffice',
                     # '--headless',
                     '--convert-to',
                     'pdf',
                     '--outdir',
                     path,
                     doc_path])
    return doc_path

def generacion_reporte(fecha_inicio = 0, fecha_fin = 0):

    numero_ven_cd = 0
    numero_ven_jfs = 0
    numero_ven_cjs = 0
    ventanillas_desordenado = []
    turnos_por_ventanilla = {}
    provincia = current_user.provincia
     
    # Obtener los datos del formulario
    fecha_actual = datetime.now()
    fecha = fecha_actual.strftime('%Y-%m-%d %H:%M:%S')
    Usuarios_ventanilla_total = Usuario.query.filter_by(rol="Ventanilla").all()
    for usuario_ventanilla in Usuarios_ventanilla_total:
        puesto = usuario_ventanilla.puesto
        ventanillas_desordenado.append(puesto)
        if puesto[0:3] == "VCD":
            numero_ven_cd += 1
            turnos_por_ventanilla[puesto] = 0
        elif puesto[0:3] == "VJF":
            numero_ven_jfs += 1
            turnos_por_ventanilla[puesto] = 0
        else:
            numero_ven_cjs += 1
            turnos_por_ventanilla[puesto] = 0 

    # Cargar la plantilla de Word
    if fecha_inicio == 0 or fecha_fin == 0: 
        doc = DocxTemplate('/media/admindpp/INFO/apps/Modelo_reportes/Modelo_reporte.docx')

        # Renderizar la plantilla con los datos del formulario
        turnos = Turnos.query.all()
        total_turnos = Turnos.query.count()
        total_turnos_cd = 0
        total_turnos_jfs = 0
        total_turnos_cjs = 0

        for turno in turnos:
            turnos_por_ventanilla[turno.puesto] += 1
            if turno.id_servicio == 1:
                total_turnos_cd += 1
            elif turno.id_servicio == 2:
                total_turnos_jfs += 1
            else:
                total_turnos_cjs += 1
        
        turnos_ventanilla_sort = {k: turnos_por_ventanilla[k] for k in sorted(turnos_por_ventanilla)}

        context = {
            'fecha': fecha,
            'provincia': provincia,
            'total_turnos': total_turnos,
            'total_turnos_cd' : total_turnos_cd,
            'total_turnos_jfs' : total_turnos_jfs,
            'total_turnos_cjs' : total_turnos_cjs,
            'turnos_por_ventanilla' : turnos_ventanilla_sort
        }
    else:
        doc = DocxTemplate('/media/admindpp/INFO/apps/Modelo_reportes/Modelo_reporte_personalizado.docx')
        # Renderizar la plantilla con los datos del formulario
        turnos = Turnos.query.filter(Turnos.fecha >= fecha_inicio, Turnos.fecha <= fecha_fin).all()
        total_turnos = Turnos.query.filter(Turnos.fecha >= fecha_inicio, Turnos.fecha <= fecha_fin).count()
        total_turnos_cd = 0
        total_turnos_jfs = 0
        total_turnos_cjs = 0
        
        for turno in turnos:
            turnos_por_ventanilla[turno.puesto] += 1
            if turno.id_servicio == 1:
                total_turnos_cd += 1
            elif turno.id_servicio == 2:
                total_turnos_jfs += 1
            else:
                total_turnos_cjs += 1
        
        turnos_ventanilla_sort = {k: turnos_por_ventanilla[k] for k in sorted(turnos_por_ventanilla)}                

        context = {
            'fecha': fecha,
            'provincia': provincia,
            'fecha_inicio': fecha_inicio,
            'fecha_fin': fecha_fin,
            'total_turnos': total_turnos,
            'total_turnos_cd' : total_turnos_cd,
            'total_turnos_jfs' : total_turnos_jfs,
            'total_turnos_cjs' : total_turnos_cjs,
            'turnos_por_ventanilla' : turnos_ventanilla_sort
        }


    servicios = ['Cambios domicilio', 'Justificaciones', 'Cajas']
    servicios_valores = [total_turnos_cd, total_turnos_jfs, total_turnos_cjs]
    colores = ['#136CB2', '#17D3E3', '#1DEEC8']
    # Crear un gráfico utilizando matplotlib
    plt.bar(servicios, servicios_valores, color=colores)
    plt.xlabel('Servicios', fontweight='bold')
    plt.ylabel('Cantidad Turnos', fontweight='bold')
    plt.title('Turnos por servicio', fontweight='bold')

    # Personalizar el estilo de las barras
    plt.gca().spines['top'].set_visible(False)  # Ocultar borde superior
    plt.gca().spines['right'].set_visible(False)  # Ocultar borde derecho
    plt.gca().tick_params(axis='x', which='both', bottom=False)  # Ocultar marcas en el eje x
    plt.gca().tick_params(axis='y', which='both', left=False)  # Ocultar marcas en el eje y
    plt.grid(axis='y', linestyle='--', alpha=0.7)  # Agregar líneas de cuadrícula horizontales

    # Guardar el gráfico como una imagen
    plt.savefig('grafico_servicios.png')
    plt.clf()
    plt.close()

    colores_puestos = ['#F1F139', '#108AF0', '#F53131', '#108AF0', '#F1F139']
    # Crear un gráfico utilizando matplotlib
    puestos = []
    puestos_valores = []
    for key, value in turnos_ventanilla_sort.items():
        puestos.append(key)
        puestos_valores.append(value)
    plt.bar(puestos, puestos_valores, color=colores_puestos)
    plt.xlabel('Ventanillas', fontweight='bold')
    plt.ylabel('Cantidad Turnos', fontweight='bold')
    plt.title('Turnos por ventanilla', fontweight='bold')

    # Personalizar el estilo de las barras
    plt.gca().spines['top'].set_visible(False)  # Ocultar borde superior
    plt.gca().spines['right'].set_visible(False)  # Ocultar borde derecho
    plt.gca().tick_params(axis='x', which='both', bottom=False)  # Ocultar marcas en el eje x
    plt.gca().tick_params(axis='y', which='both', left=False)  # Ocultar marcas en el eje y
    plt.grid(axis='y', linestyle='--', alpha=0.7)  # Agregar líneas de cuadrícula horizontales

    # # Guardar el gráfico como una imagen
    plt.savefig('grafico_puestos.png')
    plt.clf()
    plt.close()

    # Insercion de graficos al documento
    doc.render(context)
    doc.add_picture('grafico_servicios.png', width=Inches(6))
    doc.add_picture('grafico_puestos.png', width=Inches(6))

    # Guardar el nuevo documento generado
    doc.save('/media/admindpp/INFO/apps/Turnero_CNE/website/static/output.docx')

    # Ruta al documento DOCX de entrada
    docx_file = "/media/admindpp/INFO/apps/Turnero_CNE/website/static/output.docx"

    # Ruta de salida para el PDF convertido
    pdf_file = "/media/admindpp/INFO/apps/Turnero_CNE/website/static"

    try:
        doc_path = generate_pdf(docx_file, pdf_file)
        return 200
    
    except FileNotFoundError as e:
        return 500


def siguiente_id_disponible():
    max_id = db.session.query(db.func.max(Usuario.id)).scalar()
    return max_id + 1 if max_id is not None else 1

def generacion_reporte_usuario(fecha_inicio = 0, fecha_fin = 0, rol = "", id_user = 0):
    # Obtener los datos del formulario
    print(rol)
    fecha_actual = datetime.now()
    fecha = fecha_actual.strftime('%Y-%m-%d %H:%M:%S')
    user = Usuario.query.filter_by(id = id_user).first()
    servicio = user.servicio
    nombre = user.nombre
    apellido = user.apellido
    provincia = current_user.provincia
           
    if fecha_inicio == 0 or fecha_fin == 0: 
        doc = DocxTemplate('/media/admindpp/INFO/apps/Modelo_reportes/Modelo_reporte_usuario.docx')

        puesto = user.puesto
        calificaciones = Calificacion.query.filter_by(ventanilla = puesto)
        turnos = Turnos.query.all()
        total_turnos = 0
        total_turnos_cd = 0
        total_turnos_jfs = 0
        total_turnos_cjs = 0
        excelente = 0
        regular = 0
        malo = 0

        for turno in turnos:
            if turno.puesto == rol:
                total_turnos += 1
                if turno.id_servicio == 1:
                    total_turnos_cd += 1
                elif turno.id_servicio == 2:
                    total_turnos_jfs += 1
                else:
                    total_turnos_cjs += 1

        
        for calificacion in calificaciones:
            if calificacion.calificacion == "Malo":
                malo += 1
            elif calificacion.calificacion == "Regular":
                regular += 1
            else:
                excelente += 1

        context = {
            'puesto': rol,
            'fecha': fecha,
            'provincia' : provincia,
            'nombre': nombre,
            'apellido': apellido,
            'servicio' : servicio,
            'total_turnos': total_turnos,
            'total_turnos_cd' : total_turnos_cd,
            'total_turnos_jfs' : total_turnos_jfs,
            'total_turnos_cjs' : total_turnos_cjs,
            'excelente' : excelente,
            'regular' : regular,
            'malo' : malo
        }
        
    else:
        doc = DocxTemplate('/media/admindpp/INFO/apps/Modelo_reportes/Modelo_reporte_usuario_personalizado.docx')
        
        puesto = user.puesto
        calificaciones_total = Calificacion.query.filter(Calificacion.fecha >= fecha_inicio, Calificacion.fecha <= fecha_fin).all()
        turnos = Turnos.query.filter(Turnos.fecha >= fecha_inicio, Turnos.fecha <= fecha_fin).all()
        total_turnos = 0
        total_turnos_cd = 0
        total_turnos_jfs = 0
        total_turnos_cjs = 0
        calificaciones = []
        excelente = 0
        regular = 0
        malo = 0

        for turno in turnos:
            if turno.puesto == rol:
                total_turnos += 1
                if turno.id_servicio == 1:
                    total_turnos_cd += 1
                elif turno.id_servicio == 2:
                    total_turnos_jfs += 1
                else:
                    total_turnos_cjs += 1

        for calificacion in calificaciones_total:
            if calificacion.ventanilla == puesto:
                calificaciones.append(calificacion)
        
        for cal in calificaciones:
            if cal.calificacion == "Malo":
                malo += 1
            elif cal.calificacion == "Regular":
                regular += 1
            else:
                excelente += 1

            
        context = {
            'puesto': rol,
            'fecha': fecha,
            'nombre': nombre,
            'provincia': provincia,
            'apellido': apellido,
            'servicio' : servicio,
            'fecha_inicio': fecha_inicio,
            'fecha_fin': fecha_fin,
            'total_turnos': total_turnos,
            'total_turnos_cd' : total_turnos_cd,
            'total_turnos_jfs' : total_turnos_jfs,
            'total_turnos_cjs' : total_turnos_cjs,
            'excelente' : excelente,
            'regular' : regular,
            'malo' : malo
        }

    satisfaccion = ['Malo', 'Regular', 'Excelente']
    satisfaccion_valores = [malo, regular, excelente]
    colores = ['#FF0000', '#FEBB00', '#008000']
    # Crear un gráfico utilizando matplotlib
    plt.bar(satisfaccion, satisfaccion_valores, color=colores)
    plt.xlabel('Satisfaccion', fontweight='bold')
    plt.ylabel('Cantidad', fontweight='bold')
    plt.title('Satisfaccion del cliente', fontweight='bold')

    # Personalizar el estilo de las barras
    plt.gca().spines['top'].set_visible(False)  # Ocultar borde superior
    plt.gca().spines['right'].set_visible(False)  # Ocultar borde derecho
    plt.gca().tick_params(axis='x', which='both', bottom=False)  # Ocultar marcas en el eje x
    plt.gca().tick_params(axis='y', which='both', left=False)  # Ocultar marcas en el eje y
    plt.grid(axis='y', linestyle='--', alpha=0.7)  # Agregar líneas de cuadrícula horizontales

    # Guardar el gráfico como una imagen
    plt.savefig('grafico_satisfaccion.png')
    plt.clf()
    plt.close()


    # Insercion de graficos al documento
    doc.render(context)
    doc.add_picture('grafico_satisfaccion.png', width=Inches(6))

    # Guardar el nuevo documento generado
    doc.save('/media/admindpp/INFO/apps/Turnero_CNE/website/static/output_user.docx')

    # Ruta al documento DOCX de entrada
    docx_file = "/media/admindpp/INFO/apps/Turnero_CNE/website/static/output_user.docx"

    # Ruta de salida para el PDF convertido
    pdf_file = "/media/admindpp/INFO/apps/Turnero_CNE/website/static/"

    try:
        doc_path = generate_pdf(docx_file, pdf_file)
        return doc_path
    
    except FileNotFoundError as e:
        return e, 500


def calificacion_atencion(puesto, calificacion):
    
    ahora = datetime.now()
    fecha_hora_actual = ahora.strftime("%Y-%m-%d %H:%M:%S")

    ventanilla = puesto

    nueva_calificacion = Calificacion(ventanilla = ventanilla, calificacion = calificacion, fecha = fecha_hora_actual)

    db.session.add(nueva_calificacion)
    db.session.commit()
    print("Calificado correctamente")
    return 200